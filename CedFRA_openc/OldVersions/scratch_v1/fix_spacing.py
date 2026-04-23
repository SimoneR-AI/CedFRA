import docx
import os

APP_DIR = "d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav"
TEMPLATE_PATH = f"{APP_DIR}/app/templates/template_pagamento.docx"

doc = docx.Document(TEMPLATE_PATH)

# Remove "ALLEGATO "
for p in doc.paragraphs:
    if "ALLEGATO" in p.text:
        p._element.getparent().remove(p._element)
        print("Removed 'ALLEGATO'.")
        break

# Let's count empty paragraphs before the table (between title and text)
# We know the text starts with "Vogliate cortesemente"
# Let's remove up to 2 empty paragraphs before it to pull the content up
empty_paragraphs_removed = 0
for p in doc.paragraphs:
    # If we encounter the text, we stop
    if "Vogliate cortesemente" in p.text:
        break
    
    # It's an empty paragraph between title and "Vogliate..."
    if not p.text.strip() and "MODULO DI PAGAMENTO" not in p.text:
        if empty_paragraphs_removed < 3: # remove up to 3 empty lines 
            p._element.getparent().remove(p._element)
            empty_paragraphs_removed += 1

print(f"Removed {empty_paragraphs_removed} empty lines to pull content upwards.")

doc.save(TEMPLATE_PATH)
