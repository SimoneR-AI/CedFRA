import docx
import os
import shutil
from docx.oxml import OxmlElement

def insert_row_after(row):
    """Insert a new row after the given row."""
    # Create new row element
    tr = row._tr
    new_tr = OxmlElement("w:tr")
    # Copy cells from existing row to maintain formatting/borders roughly
    for tc in tr.xpath("w:tc"):
        new_tc = OxmlElement("w:tc")
        tcPr = tc.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr")
        if tcPr is not None:
            new_tc.append(import_element(tcPr))
        # Add an empty paragraph to the cell
        p = OxmlElement("w:p")
        new_tc.append(p)
        new_tr.append(new_tc)
    
    tr.addnext(new_tr)
    return docx.table._Row(new_tr, row._parent)

def import_element(element):
    """Deep copy an lxml element"""
    return type(element).fromstring(docx.oxml.xmlchemy.serialize_for_reading(element))

APP_DIR = "d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav"
TEMPLATE_PATH = os.path.join(APP_DIR, "app/templates/template_pagamento.docx")
BACKUP_PATH = os.path.join(APP_DIR, "app/templates/template_pagamento_backup.docx")

if not os.path.exists(BACKUP_PATH):
    shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)

doc = docx.Document(TEMPLATE_PATH)

print("Starting edits...")

# 1. Update Title
# Paragraph 1 is "RICHIESTA DI PAGAMENTO (“RdP”)"
title_p = doc.paragraphs[1]
# Clear all runs and add a new one with the correct text. Keep paragraph style.
title_p.clear()
run = title_p.add_run("MODULO DI PAGAMENTO (\"MdP\")")

# 2. Add NOME BANCA: {{BANCA}}
# We know it comes after IMPORTO DEL PAGAMENTO which is paragraph 12.
for i, p in enumerate(doc.paragraphs):
    if "IMPORTO DEL PAGAMENTO" in p.text:
        # Insert after this paragraph
        new_p = p.insert_paragraph_before("") # Insert before 13
        # We need to swap them essentially, but wait: `p.insert_paragraph_before` inserts BEFORE 'p'.
        # Since we found it at 'i', we want to insert before 'i+1'.
        new_p = doc.paragraphs[i+1].insert_paragraph_before("")
        r1 = new_p.add_run("NOME BANCA: ")
        r1.bold = True
        r2 = new_p.add_run("{{BANCA}}")
        r2.bold = False
        break

# 3. Update Disclaimer
# P[21] (approx)
for i, p in enumerate(doc.paragraphs):
    if "Si dichiara sotto la propria responsabilit" in p.text:
        # The text is split across runs, or in one run.
        # Let's replace text inside runs
        full_text = "".join(r.text for r in p.runs)
        
        # Replace the target text
        # Because of quotes (“ ” vs " "), we'll be careful
        target = "Procedura pagamenti terzi (“041-021” e successive Revisioni)"
        replacement = "Global Procedure 038 “Pagamenti Terzi”"
        
        if target in full_text:
            new_text = full_text.replace(target, replacement)
            # Apply back preserving style
            first_run = p.runs[0]
            first_run.text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            # Maybe the quotes are different, let's just do a naive replace
            pass
        break

# 4. Modify Table 0
t = doc.tables[0]
# Current Row 0:
# [0,0] "Prot. n. {{PROT_NUM}}\nData:"
# [0,1] "MI-PER\n{{PROT_DATA}}"

# We will create a new row using the python-docx api correctly.
# `table.add_row()` adds at the end.
new_row = t.add_row()
# We want to shift row 4->5, 3->4, 2->3, 1->2
for i in range(len(t.rows)-2, 0, -1): # from 3 down to 1
    # copy text from i to i+1
    t.cell(i+1, 0).text = t.cell(i, 0).text
    t.cell(i+1, 1).text = t.cell(i, 1).text

# Now row 1 is available for Data.
t.cell(1, 0).text = "Data (gg/mm/aaaa):"
t.cell(1, 1).text = "{{PROT_DATA}}"

# Now update row 0
t.cell(0, 0).text = "Prot. n."
t.cell(0, 1).text = "{{PROT_NUM}}"

doc.save(TEMPLATE_PATH)
print("Edits saved.")

