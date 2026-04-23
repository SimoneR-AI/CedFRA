import docx
from docx.oxml.ns import qn

APP_DIR = "d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav"
TEMPLATE_PATH = f"{APP_DIR}/app/templates/template_pagamento.docx"

doc = docx.Document(TEMPLATE_PATH)
t = doc.tables[0]

for i, row in enumerate(t.rows):
    print(f"Row {i} cells: {[c.text.strip() for c in row.cells]}")
    # Check if there is trPr -> trHeight
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is not None:
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is not None:
            val = trHeight.get(qn("w:val"))
            print(f"Row {i} has height: {val}")

    # Check paragraphs in each cell
    for j, cell in enumerate(row.cells):
        print(f"  Cell {j} paragraphs count: {len(cell.paragraphs)}")
