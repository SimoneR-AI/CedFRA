import docx
import sys

def dump_docx(filepath):
    print("====================================")
    print("FILE:", filepath)
    print("====================================")
    doc = docx.Document(filepath)
    print("PARAGRAPHS:")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"[{i}]: {text}")
    print("\nTABLES:")
    for t_idx, t in enumerate(doc.tables):
        print(f"--- TABLE {t_idx} ---")
        for r_idx, r in enumerate(t.rows):
            print(f"  ROW {r_idx}: {[c.text.strip() for c in r.cells]}")

dump_docx("d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav/newModGruppo/GP-038_All2_Modulo di pagamento.docx")
dump_docx("d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav/app/templates/template_pagamento.docx")
# Dump the "Mese Anno" file for reference of the Banca line
dump_docx("d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav/202601_JOUNEAU_GENNAIO-26_new ci.docx")
