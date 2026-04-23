import docx
import os

APP_DIR = "d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav"
TEMPLATE_PATH = os.path.join(APP_DIR, "app/templates/template_pagamento.docx")

doc = docx.Document(TEMPLATE_PATH)

# Look for the empty paragraph after DIVISA DEL PAGAMENTO
for i, p in enumerate(doc.paragraphs):
    if "DIVISA DEL PAGAMENTO" in p.text:
        # Check if the next paragraph is empty
        next_p = doc.paragraphs[i+1]
        if not next_p.text.strip():
            # Remove it
            p_xml = next_p._element
            p_xml.getparent().remove(p_xml)
            print("Removed empty paragraph.")
            break

doc.save(TEMPLATE_PATH)
