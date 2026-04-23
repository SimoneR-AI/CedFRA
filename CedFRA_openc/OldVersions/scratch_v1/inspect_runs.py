import docx
doc = docx.Document("d:/Users/Robbiani Simone/AntiG/CedFRA_antigrav/app/templates/template_pagamento.docx")

print("--- PARAGRAPHS AND RUNS ---")
for p_idx, p in enumerate(doc.paragraphs):
    print(f"P[{p_idx}]: '{p.text}'")
    for r_idx, r in enumerate(p.runs):
        print(f"  R[{r_idx}]: '{r.text}' (bold={r.bold}, italic={r.italic}, underline={r.underline})")

print("\n--- TABLE 0 RUNS ---")
t = doc.tables[0]
for r_idx, row in enumerate(t.rows):
    for c_idx, cell in enumerate(row.cells):
        print(f"CELL({r_idx},{c_idx}): '{cell.text}'")
        for p_idx, p in enumerate(cell.paragraphs):
            for run_idx, r in enumerate(p.runs):
                print(f"  R[{run_idx}]: '{r.text}'")
