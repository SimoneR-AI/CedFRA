"""Modulo per la generazione dei documenti Word da template."""

from docx import Document

from constants import MESI_ABBREVIATI, MESI_ITALIANO_UPPER


def _sostituisci_testo_run(paragraph, placeholder, nuovo_valore):
    """Sostituisce un placeholder nel testo dei run di un paragrafo.

    Cerca prima se il placeholder è interamente contenuto in un singolo run.
    Se sì, mantiene la formattazione del run. Altrimenti unisce e perde
    la formattazione (fallback di sicurezza).
    """
    # Prova a sostituire nel singolo run (preserva stili separati nella stessa riga)
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, nuovo_valore)
            return True

    # Fallback: il placeholder è spezzato su più run
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False

    new_text = full_text.replace(placeholder, nuovo_valore, 1)

    # Assegna il nuovo testo al primo run e rimuove fisicamente gli altri
    # dal DOM per evitare run vuoti che sporcano la formattazione.
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in list(paragraph.runs)[1:]:
            run._element.getparent().remove(run._element)
    return True


def _sostituisci_in_tabella(table, placeholder, nuovo_valore):
    """Sostituisce un placeholder in tutte le celle di una tabella."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _sostituisci_testo_run(paragraph, placeholder, nuovo_valore)


def _sostituisci_in_documento(doc, placeholder, nuovo_valore):
    """Sostituisce un placeholder ovunque nel documento (paragrafi + tabelle)."""
    for paragraph in doc.paragraphs:
        _sostituisci_testo_run(paragraph, placeholder, nuovo_valore)
    for table in doc.tables:
        _sostituisci_in_tabella(table, placeholder, nuovo_valore)


def genera_pagamento(template_path: str, output_path: str, dati: dict):
    """Genera il documento 'Richiesta di Pagamento' da template.

    Args:
        template_path: percorso al template_pagamento.docx
        output_path: percorso del file di output
        dati: dizionario con le chiavi:
            - prot_num: str (es. "36/2026")
            - prot_data: str (es. "18/03/2026")
            - nome_oggetto: str (es. "Redi Marchetto")
            - mese_italiano: str (es. "marzo 2026")
            - beneficiario: str
            - indirizzo: str
            - importo_completo: str (es. "€ 5.091,21 (cinquemila novantuno/21)")
            - iban: str
            - swift: str
            - data_valuta: str (es. "27 marzo 2026")
    """
    doc = Document(template_path)

    sostituzioni = {
        "{{PROT_NUM}}": dati["prot_num"],
        "{{PROT_DATA}}": dati["prot_data"],
        "{{NOME_OGGETTO}}": dati["nome_oggetto"],
        "{{MESE_ITALIANO}}": dati["mese_italiano"],
        "{{BENEFICIARIO}}": dati["beneficiario"],
        "{{INDIRIZZO}}": dati["indirizzo"],
        "{{IMPORTO_COMPLETO}}": dati["importo_completo"],
        "{{IBAN}}": dati["iban"],
        "{{SWIFT}}": dati["swift"],
        "{{BANCA}}": dati.get("banca", ""),
        "{{DATA_VALUTA}}": dati["data_valuta"],
    }

    for placeholder, valore in sostituzioni.items():
        _sostituisci_in_documento(doc, placeholder, valore)

    doc.save(output_path)


def genera_mese(template_path: str, output_path: str, dati: dict):
    """Genera il documento 'Mese Anno' (tabella riepilogo bonifico) da template.

    Args:
        template_path: percorso al template_mese.docx
        output_path: percorso del file di output
        dati: dizionario con le chiavi:
            - beneficiario: str
            - indirizzo: str
            - iban: str
            - swift: str
            - banca: str
            - importo_mese: str (es. "€ 5.091,21")
            - causale: str (es. "Marine Interiors S.p.A. – Bulletin de Salaire mars 2026")
    """
    doc = Document(template_path)

    sostituzioni = {
        "{{BENEFICIARIO}}": dati["beneficiario"],
        "{{INDIRIZZO}}": dati["indirizzo"],
        "{{IBAN}}": dati["iban"],
        "{{SWIFT}}": dati["swift"],
        "{{BANCA}}": dati["banca"],
        "{{IMPORTO_MESE}}": dati["importo_mese"],
        "{{CAUSALE}}": dati["causale"],
    }

    for placeholder, valore in sostituzioni.items():
        _sostituisci_in_documento(doc, placeholder, valore)

    doc.save(output_path)


def costruisci_nome_file_pagamento(cognome: str, mese: int, anno: int) -> str:
    """Costruisce il nome file per il documento Pagamento.

    Esempio: 202603_MARCHETTO_Pagamento_MAR26.docx
    """
    anno_breve = str(anno)[-2:]
    mese_abbr = MESI_ABBREVIATI[mese - 1]
    return f"{anno}{mese:02d}_{cognome}_Pagamento_{mese_abbr}{anno_breve}.docx"


def costruisci_nome_file_mese(cognome: str, mese: int, anno: int) -> str:
    """Costruisce il nome file per il documento Mese Anno.

    Esempio: 202603_MARCHETTO_MARZO 2026.docx
    """
    mese_upper = MESI_ITALIANO_UPPER[mese - 1]
    return f"{anno}{mese:02d}_{cognome}_{mese_upper} {anno}.docx"
